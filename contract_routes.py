from flask import Blueprint, render_template, request, flash, redirect, url_for, jsonify, current_app, send_file
from flask_login import current_user, login_required
from flask_weasyprint import HTML, CSS # For PDF generation
from docx import Document # For DOCX generation
from io import BytesIO # For handling byte streams
from models import db, Contract, User 
import datetime
import json 
import re # For cleaning HTML for text extraction

# Define contract_types_data here or import from a config file
# This data will be used to populate the UI and validate contract types
CONTRACT_TYPES_DATA = {
    "is_sozlesmesi": {
        "name": "İş Sözleşmesi",
        "category": "İş Sözleşmeleri",
        "fields": [
            {"name": "sirket_adi", "label": "Şirket Adı", "type": "text", "required": True},
            {"name": "calisan_adi", "label": "Çalışan Adı", "type": "text", "required": True},
            {"name": "pozisyon", "label": "Pozisyon", "type": "text", "required": True},
            {"name": "maas", "label": "Maaş", "type": "number", "required": True},
            {"name": "baslangic_tarihi", "label": "Başlangıç Tarihi", "type": "date", "required": True},
            {"name": "sure", "label": "Süre (Belirli/Belirsiz)", "type": "text", "required": False, "placeholder": "Belirsiz süreli"},
        ]
    },
    "gizlilik_sozlesmesi": {
        "name": "Gizlilik Sözleşmesi",
        "category": "İş Sözleşmeleri",
        "fields": [
            {"name": "taraflar", "label": "Taraflar", "type": "textarea", "required": True, "placeholder": "Örn: ABC A.Ş. ve XYZ Ltd. Şti."},
            {"name": "amac", "label": "Sözleşmenin Amacı", "type": "textarea", "required": True},
            {"name": "gizli_bilgi_tanimi", "label": "Gizli Bilgi Tanımı", "type": "textarea", "required": True},
            {"name": "sure", "label": "Süre", "type": "text", "required": True, "placeholder": "Örn: 2 yıl"},
        ]
    },
    "freelance_is_sozlesmesi": {
        "name": "Freelance İş Sözleşmesi",
        "category": "İş Sözleşmeleri",
        "fields": [
            {"name": "isveren", "label": "İşveren", "type": "text", "required": True},
            {"name": "freelancer", "label": "Freelancer", "type": "text", "required": True},
            {"name": "is_tanimi", "label": "İş Tanımı", "type": "textarea", "required": True},
            {"name": "ucret", "label": "Ücret", "type": "number", "required": True},
            {"name": "teslim_suresi", "label": "Teslim Süresi", "type": "text", "required": True, "placeholder": "Örn: 30 gün"},
        ]
    },
    "staj_sozlesmesi": {
        "name": "Staj Sözleşmesi",
        "category": "İş Sözleşmeleri",
        "fields": [
            {"name": "sirket_adi", "label": "Şirket Adı", "type": "text", "required": True},
            {"name": "stajyer_adi", "label": "Stajyer Adı", "type": "text", "required": True},
            {"name": "staj_suresi", "label": "Staj Süresi", "type": "text", "required": True, "placeholder": "Örn: 3 ay"},
            {"name": "calisma_saatleri", "label": "Çalışma Saatleri", "type": "text", "required": False, "placeholder": "Örn: 09:00 - 17:00"},
            {"name": "odeme_varsa", "label": "Ödeme (Varsa)", "type": "text", "required": False, "placeholder": "Yok"},
        ]
    },
    "uzaktan_calisma_sozlesmesi": {
        "name": "Uzaktan Çalışma Sözleşmesi",
        "category": "İş Sözleşmeleri",
        "fields": [
            {"name": "sirket_adi", "label": "Şirket Adı", "type": "text", "required": True},
            {"name": "calisan_adi", "label": "Çalışan Adı", "type": "text", "required": True},
            {"name": "pozisyon", "label": "Pozisyon", "type": "text", "required": True},
            {"name": "maas", "label": "Maaş", "type": "number", "required": True},
            {"name": "calisma_yeri", "label": "Çalışma Yeri", "type": "text", "required": True, "placeholder": "Evden"},
        ]
    },
    "konut_kira_sozlesmesi": {
        "name": "Konut Kira Sözleşmesi",
        "category": "Gayrimenkul Sözleşmeleri",
        "fields": [
            {"name": "kiraci", "label": "Kiracı", "type": "text", "required": True},
            {"name": "kiraya_veren", "label": "Kiraya Veren", "type": "text", "required": True},
            {"name": "mulk_adresi", "label": "Mülk Adresi", "type": "textarea", "required": True},
            {"name": "kira_miktari", "label": "Kira Miktarı", "type": "number", "required": True},
            {"name": "depozito_miktari", "label": "Depozito Miktarı", "type": "number", "required": False},
            {"name": "kira_suresi", "label": "Kira Süresi", "type": "text", "required": True, "placeholder": "Örn: 1 yıl"},
            {"name": "diger_kosullar", "label": "Diğer Koşullar", "type": "textarea", "required": False, "placeholder": "Aidat, bakım vb."},
        ]
    },
    "isyeri_kira_sozlesmesi": {
        "name": "İşyeri Kira Sözleşmesi",
        "category": "Gayrimenkul Sözleşmeleri",
        "fields": [
            {"name": "kiraci_unvan", "label": "Kiracı (Ünvan/Ad Soyad)", "type": "text", "required": True},
            {"name": "kiraya_veren_unvan", "label": "Kiraya Veren (Ünvan/Ad Soyad)", "type": "text", "required": True},
            {"name": "isyeri_adresi", "label": "İşyeri Adresi", "type": "textarea", "required": True},
            {"name": "kira_bedeli", "label": "Kira Bedeli (Aylık)", "type": "number", "required": True},
            {"name": "stopaj_durumu", "label": "Stopaj Durumu", "type": "text", "required": False, "placeholder": "Kiracı tarafından ödenecektir"},
            {"name": "kira_suresi", "label": "Kira Süresi", "type": "text", "required": True, "placeholder": "Örn: 3 yıl"},
            {"name": "depozito", "label": "Depozito", "type": "number", "required": False},
        ]
    },
    "tasinmaz_satis_vaadi_sozlesmesi": {
        "name": "Taşınmaz Satış Vaadi Sözleşmesi",
        "category": "Gayrimenkul Sözleşmeleri",
        "fields": [
            {"name": "vaad_eden", "label": "Satış Vaadinde Bulunan", "type": "text", "required": True},
            {"name": "vaad_alan", "label": "Lehine Satış Vaadinde Bulunulan", "type": "text", "required": True},
            {"name": "tasinmaz_bilgileri", "label": "Taşınmaz Bilgileri (Ada, Parsel, Adres)", "type": "textarea", "required": True},
            {"name": "satis_bedeli", "label": "Satış Bedeli", "type": "number", "required": True},
            {"name": "odeme_sekli", "label": "Ödeme Şekli", "type": "textarea", "required": True},
            {"name": "teslim_tarihi", "label": "Teslim Tarihi", "type": "date", "required": False},
        ]
    },
    "temizlik_hizmet_sozlesmesi": {
        "name": "Temizlik Hizmet Sözleşmesi",
        "category": "Hizmet Sözleşmeleri",
        "fields": [
            {"name": "hizmet_alan", "label": "Hizmet Alan (Müşteri)", "type": "text", "required": True},
            {"name": "hizmet_veren", "label": "Hizmet Veren (Firma)", "type": "text", "required": True},
            {"name": "hizmet_kapsami", "label": "Hizmet Kapsamı", "type": "textarea", "required": True, "placeholder": "Ofis genel temizliği, cam silme vb."},
            {"name": "hizmet_periyodu", "label": "Hizmet Periyodu", "type": "text", "required": True, "placeholder": "Haftada 2 kez, Aylık vb."},
            {"name": "ucret", "label": "Ücret", "type": "number", "required": True},
            {"name": "malzemeler", "label": "Malzemeler", "type": "text", "required": False, "placeholder": "Hizmet veren tarafından temin edilecektir"},
        ]
    },
    "avukatlik_sozlesmesi": {
        "name": "Avukatlık Sözleşmesi",
        "category": "Hizmet Sözleşmeleri",
        "fields": [
            {"name": "muvekkil", "label": "Müvekkil", "type": "text", "required": True},
            {"name": "avukat", "label": "Avukat/Avukatlık Bürosu", "type": "text", "required": True},
            {"name": "is_konusu", "label": "İşin Konusu (Dava/Danışmanlık)", "type": "textarea", "required": True},
            {"name": "vekalet_ucreti", "label": "Vekalet Ücreti", "type": "text", "required": True, "placeholder": "Örn: 10.000 TL + KDV"},
            {"name": "masraflar", "label": "Masraflar", "type": "text", "required": False, "placeholder": "Müvekkil tarafından karşılanacaktır"},
        ]
    },
    "mal_alim_satim_sozlesmesi": {
        "name": "Mal Alım Satım Sözleşmesi",
        "category": "Ticari Sözleşmeler",
        "fields": [
            {"name": "satici", "label": "Satıcı", "type": "text", "required": True},
            {"name": "alici", "label": "Alıcı", "type": "text", "required": True},
            {"name": "mal_cinsi_miktari", "label": "Malın Cinsi ve Miktarı", "type": "textarea", "required": True},
            {"name": "birim_fiyat_toplam_bedel", "label": "Birim Fiyat ve Toplam Bedel", "type": "text", "required": True},
            {"name": "teslim_yeri_sekli", "label": "Teslim Yeri ve Şekli", "type": "text", "required": True},
            {"name": "odeme_vadesi", "label": "Ödeme Vadesi", "type": "text", "required": True},
        ]
    },
    "danismanlik_hizmet_sozlesmesi": {
        "name": "Danışmanlık Hizmet Sözleşmesi",
        "category": "Ticari Sözleşmeler",
        "fields": [
            {"name": "danisan", "label": "Danışan (Hizmeti Alan)", "type": "text", "required": True},
            {"name": "danisman", "label": "Danışman (Hizmeti Veren)", "type": "text", "required": True},
            {"name": "danismanlik_konusu", "label": "Danışmanlık Konusu", "type": "textarea", "required": True},
            {"name": "sure", "label": "Süre", "type": "text", "required": True, "placeholder": "Örn: 6 ay"},
            {"name": "ucret_odeme_kosullari", "label": "Ücret ve Ödeme Koşulları", "type": "textarea", "required": True},
        ]
    },
    "genel_amacli_sozlesme": {
        "name": "Genel Amaçlı Sözleşme Taslağı",
        "category": "Genel Sözleşme",
        "fields": [
            {"name": "taraf_1", "label": "Taraflardan Biri (Adı/Ünvanı)", "type": "text", "required": True},
            {"name": "taraf_2", "label": "Diğer Taraf (Adı/Ünvanı)", "type": "text", "required": True},
            {"name": "sozlesme_konusu", "label": "Sözleşmenin Konusu", "type": "textarea", "required": True},
            {"name": "temel_sartlar", "label": "Temel Şartlar ve Yükümlülükler", "type": "textarea", "required": True},
            {"name": "sure_bedel_opsiyonel", "label": "Süre / Bedel (Varsa)", "type": "text", "required": False},
            {"name": "ek_maddeler", "label": "Ek Maddeler (Varsa)", "type": "textarea", "required": False},
        ]
    }
}

contract_bp = Blueprint('contract', __name__, url_prefix='/contract', template_folder='templates')

@contract_bp.before_request
@login_required
def require_login():
    pass

@contract_bp.route('/')
def index():
    """Displays the main page for selecting contract types."""
    print("--- contract_routes.py: index route CALLED ---") 
    print(f"DEBUG: Type of CONTRACT_TYPES_DATA: {type(CONTRACT_TYPES_DATA)}") 
    print(f"DEBUG: Length of CONTRACT_TYPES_DATA: {len(CONTRACT_TYPES_DATA) if isinstance(CONTRACT_TYPES_DATA, dict) else 'Not a dict'}") 
    
    categorized_contracts = {}
    try:
        for key, value in CONTRACT_TYPES_DATA.items():
            if not isinstance(value, dict): 
                print(f"DEBUG: Item with key '{key}' in CONTRACT_TYPES_DATA is not a dictionary. Value: {value}")
                continue 

            category = value.get("category", "Diğer")
            if category not in categorized_contracts:
                categorized_contracts[category] = []
            
            item_name = value.get("name", f"Unnamed Contract ({key})") 
            item_description = value.get("description", "")

            categorized_contracts[category].append({
                "key": key, 
                "name": item_name, 
                "description": item_description
            })
    except Exception as e:
        print(f"DEBUG: Error during categorization loop: {e}")

    print(f"DEBUG: categorized_contracts POPULATED: {json.dumps(categorized_contracts, indent=2, ensure_ascii=False)}") 
    
    user_contracts = Contract.query.filter_by(user_id=current_user.id, is_deleted=False).order_by(Contract.created_at.desc()).all()
    print(f"DEBUG: Fetched {len(user_contracts)} user_contracts.")

    print("--- contract_routes.py: index route RENDERING TEMPLATE ---")
    return render_template('contracts/contract_hub.html',
                           title='Sözleşme Hazırla',
                           categorized_contracts=categorized_contracts,
                           user_contracts=user_contracts,
                           CONTRACT_TYPES_DATA=CONTRACT_TYPES_DATA)

@contract_bp.route('/get_contract_form/<string:contract_type_key>', methods=['GET'])
def get_contract_form(contract_type_key):
    """Returns the specific form fields for a selected contract type."""
    contract_details = CONTRACT_TYPES_DATA.get(contract_type_key)
    if not contract_details:
        return jsonify({"error": "Geçersiz sözleşme türü"}), 404
    
    return jsonify({
        "name": contract_details["name"],
        "fields": contract_details["fields"]
    })

@contract_bp.route('/generate', methods=['POST'])
def generate_contract():
    """
    Handles the contract generation request.
    Receives form data, calls AI service, saves the contract.
    """
    data = request.json
    contract_type_key = data.get('contract_type')
    form_inputs = data.get('form_inputs')
    custom_title = data.get('title', None)
    custom_prompt = data.get('custom_prompt', '')

    if not contract_type_key or not form_inputs:
        return jsonify({"error": "Eksik bilgi: sözleşme türü ve form girdileri gereklidir."}), 400

    contract_template_info = CONTRACT_TYPES_DATA.get(contract_type_key)
    if not contract_template_info:
        return jsonify({"error": "Geçersiz sözleşme türü."}), 400

    from ai import generate_contract_with_ai 
    
    ai_generated_html_content, ai_generated_text_content = generate_contract_with_ai(
        contract_template_info['name'], 
        form_inputs,
        custom_prompt 
    )

    try:
        new_contract = Contract(
            user_id=current_user.id,
            contract_type=contract_type_key,
            title=custom_title if custom_title else contract_template_info['name'],
            input_data=form_inputs,
            generated_content_html=ai_generated_html_content,
            generated_content_text=ai_generated_text_content
        )
        db.session.add(new_contract)
        db.session.commit()
        flash('Sözleşmeniz başarıyla oluşturuldu!', 'success')
        return jsonify({
            "message": "Sözleşme başarıyla oluşturuldu!",
            "contract_id": new_contract.id,
            "contract_html": ai_generated_html_content
        }), 201
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Sözleşme kaydedilirken hata: {e}")
        return jsonify({"error": f"Sözleşme kaydedilirken bir hata oluştu: {str(e)}"}), 500

@contract_bp.route('/view/<int:contract_id>')
def view_contract(contract_id):
    """Displays a previously generated contract."""
    contract = Contract.query.filter_by(id=contract_id, user_id=current_user.id, is_deleted=False).first_or_404()
    return render_template('contracts/view_contract.html', 
                           title=contract.title, 
                           contract=contract,
                           CONTRACT_TYPES_DATA=CONTRACT_TYPES_DATA)

@contract_bp.route('/delete/<int:contract_id>', methods=['POST'])
def delete_contract(contract_id):
    """Soft deletes a contract."""
    contract_to_delete = Contract.query.filter_by(id=contract_id, user_id=current_user.id, is_deleted=False).first_or_404()
    try:
        contract_to_delete.is_deleted = True
        contract_to_delete.deleted_at = datetime.datetime.utcnow()
        db.session.commit()
        flash(f"'{contract_to_delete.title}' başlıklı sözleşme başarıyla silindi.", 'success')
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f"Sözleşme silinirken hata: {e}")
        flash(f"Sözleşme silinirken bir hata oluştu: {str(e)}", 'danger')
    return redirect(url_for('contract.index'))

@contract_bp.route('/export/pdf/<int:contract_id>')
@login_required
def export_pdf(contract_id):
    contract = Contract.query.filter_by(id=contract_id, user_id=current_user.id, is_deleted=False).first_or_404()
    
    html_content = contract.generated_content_html
    if not html_content:
        flash('Sözleşme içeriği bulunamadı, PDF oluşturulamıyor.', 'danger')
        return redirect(url_for('contract.view_contract', contract_id=contract_id))

    try:
        # Add some basic styling for the PDF if not already in html_content
        # You might want to link a CSS file or embed styles for better PDF appearance
        html_for_pdf = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>{contract.title}</title>
            <style>
                body {{ font-family: 'DejaVu Sans', sans-serif; line-height: 1.6; }}
                h1, h2, h3, h4, h5, h6 {{ page-break-after: avoid; }}
                p {{ margin-bottom: 0.5em; }}
                /* Add more specific styles as needed */
            </style>
        </head>
        <body>
            <h1>{contract.title}</h1>
            {html_content}
        </body>
        </html>
        """
        pdf_file = HTML(string=html_for_pdf).write_pdf()
        
        return send_file(
            BytesIO(pdf_file),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"{contract.title.replace(' ', '_').lower()}_{contract_id}.pdf"
        )
    except Exception as e:
        current_app.logger.error(f"PDF oluşturulurken hata: {e}")
        flash(f"PDF oluşturulurken bir hata oluştu: {str(e)}", 'danger')
        return redirect(url_for('contract.view_contract', contract_id=contract_id) if contract_id else url_for('contract.index'))


def clean_html_for_docx(html_content):
    """Basic HTML tag removal for DOCX conversion. More sophisticated parsing might be needed."""
    if not html_content:
        return ""
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', ' ', html_content)
    # Replace common HTML entities
    text = text.replace('&nbsp;', ' ')
    text = text.replace('&', '&')
    text = text.replace('<', '<')
    text = text.replace('>', '>')
    text = text.replace('"', '"')
    text = text.replace('&#39;', "'")
    # Collapse multiple spaces
    text = re.sub(r'\s+', ' ', text).strip()
    return text

@contract_bp.route('/edit/<int:contract_id>', methods=['GET', 'POST'])
@login_required
def edit_contract(contract_id):
    contract_to_edit = Contract.query.filter_by(id=contract_id, user_id=current_user.id, is_deleted=False).first_or_404()

    if request.method == 'POST':
        new_title = request.form.get('contract_title')
        new_html_content = request.form.get('contract_content')

        if not new_title or not new_title.strip():
            flash('Sözleşme başlığı boş bırakılamaz.', 'danger')
            # Rerender the edit page with current data if title is empty
            return render_template('contracts/edit_contract.html',
                                   title=f"Düzenle: {contract_to_edit.title}",
                                   contract=contract_to_edit,
                                   CONTRACT_TYPES_DATA=CONTRACT_TYPES_DATA,
                                   form={}) # Pass an empty dict or a proper form object if using Flask-WTF
        
        contract_to_edit.title = new_title.strip()
        contract_to_edit.generated_content_html = new_html_content
        contract_to_edit.generated_content_text = clean_html_for_docx(new_html_content) # Update plain text version
        contract_to_edit.updated_at = datetime.datetime.utcnow()
        
        try:
            db.session.commit()
            flash(f"'{contract_to_edit.title}' başlıklı sözleşme başarıyla güncellendi.", 'success')
            return redirect(url_for('contract.view_contract', contract_id=contract_to_edit.id))
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f"Sözleşme güncellenirken hata: {e}")
            flash(f"Sözleşme güncellenirken bir hata oluştu: {str(e)}", 'danger')
            # Rerender edit page with submitted data to avoid data loss on error
            # For simplicity, we're just re-fetching, but ideally, you'd pass back the submitted form data
            contract_to_edit.title = request.form.get('contract_title') # Show submitted title
            contract_to_edit.generated_content_html = request.form.get('contract_content') # Show submitted content

    return render_template('contracts/edit_contract.html', 
                           title=f"Düzenle: {contract_to_edit.title}", 
                           contract=contract_to_edit,
                           CONTRACT_TYPES_DATA=CONTRACT_TYPES_DATA,
                           form={}) # Pass an empty dict or a proper form object if using Flask-WTF

@contract_bp.route('/export/docx/<int:contract_id>')
@login_required
def export_docx(contract_id):
    contract = Contract.query.filter_by(id=contract_id, user_id=current_user.id, is_deleted=False).first_or_404()
    
    # Prefer text content if available, otherwise clean HTML
    text_content = contract.generated_content_text
    if not text_content and contract.generated_content_html:
        text_content = clean_html_for_docx(contract.generated_content_html)
    elif not text_content and not contract.generated_content_html:
        flash('Sözleşme içeriği bulunamadı, Word belgesi oluşturulamıyor.', 'danger')
        return redirect(url_for('contract.view_contract', contract_id=contract_id))

    try:
        document = Document()
        document.add_heading(contract.title, level=1)
        
        # Split content into paragraphs. This is a basic split.
        # For better formatting, you might need to parse HTML structure more carefully.
        paragraphs = text_content.split('\n')
        for para_text in paragraphs:
            if para_text.strip(): # Add non-empty paragraphs
                document.add_paragraph(para_text)
            else: # Add an empty paragraph for line breaks if desired
                document.add_paragraph()


        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0) # Reset stream position to the beginning

        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f"{contract.title.replace(' ', '_').lower()}_{contract_id}.docx"
        )
    except Exception as e:
        current_app.logger.error(f"Word belgesi oluşturulurken hata: {e}")
        flash(f"Word belgesi oluşturulurken bir hata oluştu: {str(e)}", 'danger')
        return redirect(url_for('contract.view_contract', contract_id=contract_id) if contract_id else url_for('contract.index'))
